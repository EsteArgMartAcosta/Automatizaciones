from docx import Document
from docx.shared import Pt
from pathlib import Path

# Ruta al documento original
archivo = Path.home() / "Downloads" / "Telegram" / "Toldito 08 08 25.docx"
documento = Document(archivo)

noticias = []
actual = []

# Separar noticias por títulos en negrita
for parrafo in documento.paragraphs:
    if parrafo.runs and any(run.bold for run in parrafo.runs):
        if actual:
            noticias.append(actual)
            actual = []
    actual.append(parrafo)

if actual:
    noticias.append(actual)

# Carpeta de salida
output_dir = archivo.parent / "noticias_separadas"
output_dir.mkdir(exist_ok=True)

# Guardar cada noticia
for i, bloques in enumerate(noticias, start=1):
    nuevo_doc = Document()
    
    for j, parrafo in enumerate(bloques):
        nuevo_parrafo = nuevo_doc.add_paragraph()
        
        for run in parrafo.runs:
            nuevo_run = nuevo_parrafo.add_run(run.text)
            nuevo_run.bold = run.bold
            nuevo_run.italic = run.italic
            nuevo_run.underline = run.underline
            nuevo_run.font.name = "Times New Roman"
            if run.font.size:
                nuevo_run.font.size = run.font.size

        # Si es el título
        if j == 0:
            hashtag = nuevo_parrafo.add_run(" #Tolditos7días")
            hashtag.bold = True
            hashtag.font.name = "Times New Roman"
            hashtag.font.size = Pt(28)

            for r in nuevo_parrafo.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(28)
                r.bold = True

    # Guardar archivo
    nombre_archivo = output_dir / f"noticia_{i}.docx"
    nuevo_doc.save(nombre_archivo)
    print(f"✅ {nombre_archivo.name} creada")

print(f"\n🎉 {len(noticias)} noticias separadas y guardadas en: {output_dir.resolve()}")
